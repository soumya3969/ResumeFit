"""
Text Extraction Module
Extracts text content from PDF and DOCX resume files.
"""

import fitz  # PyMuPDF
from docx import Document
from typing import Optional
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    Extract text from a PDF file using PyMuPDF.
    
    Args:
        file_path: Path to the PDF file or file-like object
        
    Returns:
        Extracted text as a string, or None if extraction fails
    """
    try:
        text = ""
        
        # Handle both file paths and file-like objects (for Streamlit)
        if isinstance(file_path, str):
            doc = fitz.open(file_path)
        else:
            # For file-like objects from Streamlit
            doc = fitz.open(stream=file_path.read(), filetype="pdf")
        
        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            text += page_text
            logger.info(f"Extracted text from page {page_num}")
        
        doc.close()
        
        if not text.strip():
            logger.warning("No text extracted from PDF")
            return None
            
        logger.info(f"Successfully extracted {len(text)} characters from PDF")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return None


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extract text from a DOCX file using python-docx.
    
    Args:
        file_path: Path to the DOCX file or file-like object
        
    Returns:
        Extracted text as a string, or None if extraction fails
    """
    try:
        # Handle both file paths and file-like objects (for Streamlit)
        if isinstance(file_path, str):
            doc = Document(file_path)
        else:
            # For file-like objects from Streamlit
            doc = Document(file_path)
        
        # Extract text from paragraphs
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        full_text = "\n".join(text)
        
        if not full_text.strip():
            logger.warning("No text extracted from DOCX")
            return None
            
        logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
        return full_text
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return None


def extract_text(file_path, file_type: str) -> Optional[str]:
    """
    Main function to extract text based on file type.
    
    Args:
        file_path: Path to the file or file-like object
        file_type: Type of file ('pdf' or 'docx')
        
    Returns:
        Extracted text as a string, or None if extraction fails
    """
    file_type = file_type.lower()
    
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type in ['docx', 'doc']:
        return extract_text_from_docx(file_path)
    else:
        logger.error(f"Unsupported file type: {file_type}")
        return None
